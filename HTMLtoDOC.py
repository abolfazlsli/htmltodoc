from bs4 import BeautifulSoup
from docx import Document

def html_to_docx(html_file, docx_file):
    
    with open(html_file, 'r', encoding='utf-8') as f:
        html_content = f.read()

    
    soup = BeautifulSoup(html_content, 'html.parser')

    
    doc = Document()

    
    for element in soup.body.find_all(['h1', 'h2', 'h3', 'p']):
        if element.name == 'h1':
            doc.add_heading(element.get_text(), level=1)
        elif element.name == 'h2':
            doc.add_heading(element.get_text(), level=2)
        elif element.name == 'h3':
            doc.add_heading(element.get_text(), level=3)
        elif element.name == 'p':
            doc.add_paragraph(element.get_text())

    
    doc.save(docx_file)


html_file = 'example.html'  
docx_file = 'output.docx'   
html_to_docx(html_file, docx_file)

print("new file generated ")
